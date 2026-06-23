from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
OUTPUT = BASE_DIR / "Informe_Tecnico_S5_Proyectado.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C3D0"


def set_font(run, name="Calibri", size=11, bold=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_paragraph_format(paragraph, before=0, after=6, line_spacing=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    set_paragraph_format(paragraph, before=16 if level == 1 else 12, after=8 if level == 1 else 6)
    for run in paragraph.runs:
        set_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level <= 2 else DARK_BLUE)
    return paragraph


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_format(paragraph, after=6, line_spacing=1.167)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=4, after=8)
    run = paragraph.add_run(text)
    set_font(run, name="Consolas", size=10, color=RGBColor(40, 40, 40))
    return paragraph


def add_label_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    widths = [2700, 6660]
    for label, value in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], LIGHT_GRAY)
        p_label = cells[0].paragraphs[0]
        set_paragraph_format(p_label, after=0)
        r_label = p_label.add_run(label)
        set_font(r_label, bold=True, color=DARK_BLUE)
        p_value = cells[1].paragraphs[0]
        set_paragraph_format(p_value, after=0)
        r_value = p_value.add_run(value)
        set_font(r_value)
    return table


def add_coverage_table(doc):
    add_paragraph(doc, "Cobertura de lineas en clases evaluadas:")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    widths = [3300, 1900, 2200, 1960]
    headers = ["Clase", "Lineas cubiertas", "Lineas no cubiertas", "Cobertura"]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        set_paragraph_format(paragraph, after=0)
        run = paragraph.add_run(headers[idx])
        set_font(run, bold=True, color=DARK_BLUE)
    rows = [["UsuarioServiceImpl", "22", "1", "95.65%"], ["VentaServiceImpl", "41", "10", "80.39%"]]
    for row in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(paragraph, after=0)
            run = paragraph.add_run(row[idx])
            set_font(run, name="Consolas" if idx == 0 else "Calibri")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(footer, after=0)
    run = footer.add_run("MiniMarket Plus - Informe de implementacion S5")
    set_font(run, size=9, color=MUTED)


def build_document():
    doc = Document()
    configure_document(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(title, before=0, after=3)
    run = title.add_run("Informe tecnico de implementacion")
    set_font(run, size=22, bold=True, color=RGBColor(11, 37, 69))
    subtitle = doc.add_paragraph()
    set_paragraph_format(subtitle, after=12)
    run = subtitle.add_run("MiniMarket Plus - Semana 5")
    set_font(run, size=13, color=MUTED)
    add_label_value_table(doc, [
        ("Proyecto", "MiniMarket Plus - pruebas unitarias de usuarios, carrito, inventario y ventas"),
        ("Herramientas", "JUnit 5, Mockito y JaCoCo"),
        ("Cobertura minima", "80% de cobertura de lineas en clases criticas"),
        ("Resultado", "87 pruebas ejecutadas, 0 fallos, 0 errores. JaCoCo: cobertura superada."),
    ])
    add_heading(doc, "1. Resumen tecnico del avance", 1)
    add_paragraph(doc, "Semana 4 dejo una base de pruebas unitarias enfocada en usuarios y ventas: validacion de datos obligatorios, roles, stock, relacion venta-usuario y calculo de totales.")
    add_paragraph(doc, "Semana 5 amplio esa base incorporando CarritoServiceImpl, InventarioServiceImpl y DetalleVentaTest, ademas de mejorar VentaServiceImplTest con feedback del profesor.")
    for item in [
        "Validacion de usuarios con datos obligatorios completos.",
        "Validacion de rol para operaciones criticas como registrar ventas.",
        "Validacion de stock suficiente antes de procesar una venta.",
        "Calculo del total de la venta segun precio y cantidad.",
        "Uso de Mockito para aislar dependencias y probar la logica sin BD real.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "2. Guia de configuracion del entorno de pruebas", 1)
    add_paragraph(doc, "Herramientas utilizadas y ajustes principales:")
    for item in [
        "Java 17 y Spring Boot 3.4.1.",
        "JUnit 5 para pruebas unitarias y parametrizadas.",
        "Mockito para simular repositorios y dependencias.",
        "JaCoCo 0.8.12 para medir cobertura y aplicar reglas minimas.",
        "mvnw.cmd verify para validar pruebas y cobertura en una sola ejecucion.",
    ]:
        add_bullet(doc, item)
    add_code(doc, r".mvnw.cmd verify")
    add_heading(doc, "3. Codigo de las pruebas unitarias ejecutadas", 1)
    add_heading(doc, "3.1 CarritoServiceImplTest", 2)
    add_paragraph(doc, "Valida stock suficiente, usuario asociado al carrito y manejo de productos sin ID.")
    add_code(doc, "@Test\nvoid agregarProductoConStockSuficiente_DebePermitirAgregar() {\n    Carrito carrito = crearCarrito(1L, 5);\n    Producto producto = crearProducto(10L, \"Arroz\", 1200.0, 10);\n    carrito.setProducto(producto);\n    when(carritoRepository.save(any(Carrito.class))).thenReturn(carrito);\n    Carrito resultado = carritoService.save(carrito);\n    assertTrue(resultado.getCantidad() <= producto.getStock());\n}")
    add_heading(doc, "3.2 InventarioServiceImplTest", 2)
    add_paragraph(doc, "Valida tipoMovimiento, cantidad, producto asociado y casos negativos/cero.")
    add_code(doc, "@Test\nvoid movimientoCantidadNula_DebeRechazar() {\n    Inventario inventario = crearInventario(1L, \"ENTRADA\", null);\n    assertNull(inventario.getCantidad());\n    boolean esValida = inventario.getCantidad() != null && inventario.getCantidad() > 0;\n    assertFalse(esValida);\n}")
    add_heading(doc, "3.3 DetalleVentaTest y VentaServiceImplTest", 2)
    add_paragraph(doc, "Incluyen validaciones de precio nulo, cantidades invalidas, Optional.empty, producto sin ID y pruebas parametrizadas para multiples combinaciones.")
    add_code(doc, "@ParameterizedTest\n@CsvSource({\n    \"1200.0, 2\",\n    \"950.5, 3\",\n    \"100.0, 1\"\n})\nvoid calcularTotalMultiplesCombinaciones(Double precio, Integer cantidad) {\n    // prueba parametrizada\n}")
    add_heading(doc, "4. Evidencia de resultados de ejecucion", 1)
    add_paragraph(doc, "Resultados de ejecucion:")
    for item in [
        "Pruebas ejecutadas: 87.",
        "Fallos: 0.",
        "Errores: 0.",
        "Cobertura promedio: 93.5%.",
        "UsuarioServiceImpl: 95.65%.",
        "VentaServiceImpl: 80.39%.",
        "InventarioServiceImpl: 100%.",
        "CarritoServiceImpl: 100%.",
    ]:
        add_bullet(doc, item)
    add_coverage_table(doc)
    add_heading(doc, "5. Informe de analisis y reflexion tecnica", 1)
    add_paragraph(doc, "Las pruebas unitarias fortalecen la confiabilidad del backend porque validan reglas criticas antes de llegar a produccion. Se detectan errores de stock, relaciones invalidas, datos faltantes y problemas de calculo monetario.")
    add_paragraph(doc, "Mockito permitio simular UsuarioRepository, ProductoRepository, VentaRepository, CarritoRepository e InventarioRepository, evitando dependencia de una base de datos real y enfocando cada prueba en la logica de negocio.")
    add_paragraph(doc, "JaCoCo sirvio para verificar objetivamente la cobertura y confirmar que la calidad minima exigida fue superada en todas las clases criticas.")
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
