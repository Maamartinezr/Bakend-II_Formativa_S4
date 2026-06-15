from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT = BASE_DIR / "INFORME_IMPLEMENTACION_S4.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C3D0"


def set_font(run, name="Calibri", size=11, bold=False, color=INK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_paragraph_format(paragraph, before=0, after=6, line_spacing=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    set_paragraph_format(paragraph, before=16 if level == 1 else 12, after=8 if level == 1 else 6)
    for run in paragraph.runs:
        set_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level <= 2 else DARK_BLUE)
    return paragraph


def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_format(paragraph, after=6, line_spacing=1.167)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, before=4, after=8)
    run = paragraph.add_run(text)
    set_font(run, name="Consolas", size=10, color=RGBColor(40, 40, 40))
    return paragraph


def add_label_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    widths = [2700, 6660]

    for label, value in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], LIGHT_GRAY)
        p_label = cells[0].paragraphs[0]
        set_paragraph_format(p_label, after=0)
        r_label = p_label.add_run(label)
        set_font(r_label, bold=True, color=DARK_BLUE)
        p_value = cells[1].paragraphs[0]
        set_paragraph_format(p_value, after=0)
        r_value = p_value.add_run(value)
        set_font(r_value)
    return table


def add_coverage_table(doc):
    add_paragraph(doc, "Cobertura de lineas en clases evaluadas:")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    widths = [3300, 1900, 2200, 1960]
    headers = ["Clase", "Lineas cubiertas", "Lineas no cubiertas", "Cobertura"]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_margins(cell)
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        set_paragraph_format(paragraph, after=0)
        run = paragraph.add_run(headers[idx])
        set_font(run, bold=True, color=DARK_BLUE)

    rows = [
        ["UsuarioServiceImpl", "22", "1", "95.65%"],
        ["VentaServiceImpl", "41", "10", "80.39%"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(paragraph, after=0)
            run = paragraph.add_run(row[idx])
            set_font(run, name="Consolas" if idx == 0 else "Calibri")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(footer, after=0)
    run = footer.add_run("MiniMarket Plus - Informe de implementacion S4")
    set_font(run, size=9, color=MUTED)


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(title, before=0, after=3)
    run = title.add_run("Informe tecnico de implementacion")
    set_font(run, size=22, bold=True, color=RGBColor(11, 37, 69))

    subtitle = doc.add_paragraph()
    set_paragraph_format(subtitle, after=12)
    run = subtitle.add_run("MiniMarket Plus - Semana 4")
    set_font(run, size=13, color=MUTED)

    add_label_value_table(
        doc,
        [
            ("Proyecto", "MiniMarket Plus - pruebas unitarias de usuarios y ventas"),
            ("Herramientas", "JUnit 5, Mockito y JaCoCo"),
            ("Cobertura minima", "80% de cobertura de lineas en UsuarioServiceImpl y VentaServiceImpl"),
            ("Resultado", "17 pruebas ejecutadas, 0 fallos, 0 errores. JaCoCo: All coverage checks have been met."),
        ],
    )

    add_heading(doc, "1. Resumen tecnico del avance", 1)
    add_paragraph(
        doc,
        "El proyecto de Semana 4 fue preparado a partir del archivo PBY2202_Exp2_S4_Caso_actividad_MiniMarketPlusTest (forma C).zip. "
        "Tambien se reviso el proyecto de Semana 3, Proyecto_minimarket_S3, porque contenia avances mas completos en seguridad, validaciones y documentacion. "
        "Para esta actividad se mantuvo como base el proyecto S4, ya que el foco solicitado corresponde a pruebas unitarias sobre funcionalidades de usuarios y ventas.",
    )
    add_paragraph(
        doc,
        "Durante Semana 3 el sistema ya contaba con una estructura base de microservicio Spring Boot: entidades JPA, repositorios, servicios, controladores REST, configuracion H2 y una base de seguridad. "
        "En Semana 4 se seleccionaron para evaluacion unitaria las reglas directamente relacionadas con el caso.",
    )
    for item in [
        "Validacion de usuarios con datos obligatorios completos.",
        "Validacion de rol para operaciones criticas como registrar ventas.",
        "Validacion de stock suficiente antes de procesar una venta.",
        "Calculo del total de la venta segun precio y cantidad de productos.",
        "Simulacion de dependencias con Mockito para probar servicios de forma aislada.",
    ]:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "Nota: el enunciado del informe menciona Reserva y Proveedor, pero el contexto, pasos y necesidades especificas de esta forma C se refieren a Usuario y Venta. "
        "Por coherencia tecnica, el informe y las pruebas se enfocan en Usuario y Venta.",
    )

    add_heading(doc, "2. Guia de configuracion", 1)
    add_paragraph(doc, "El proyecto fue extraido y trabajado en:")
    add_code(doc, r"C:\Users\Marialex\Documents\Backend II\MiniMarketPlus_S4_Testing\minimarket")
    add_paragraph(doc, "Se agregaron o confirmaron las herramientas solicitadas en pom.xml:")
    for item in [
        "JUnit 5: incluido mediante spring-boot-starter-test.",
        "Mockito: agregado explicitamente con mockito-junit-jupiter.",
        "JaCoCo: agregado con jacoco-maven-plugin version 0.8.12.",
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, "La configuracion de JaCoCo quedo con tres ejecuciones:")
    for item in [
        "prepare-agent: instrumenta el codigo durante las pruebas.",
        "report: genera el reporte HTML y CSV en target/site/jacoco.",
        "check: valida en fase verify que UsuarioServiceImpl y VentaServiceImpl cumplan al menos 80% de cobertura de lineas.",
    ]:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "Tambien se normalizo application.properties a UTF-8/ASCII, ya que Maven fallaba al copiar recursos por un problema de codificacion en comentarios con acentos.",
    )

    add_heading(doc, "3. Diseno y resultados de pruebas", 1)
    add_heading(doc, "Usuario", 2)
    add_paragraph(doc, "Archivo de pruebas: src/test/java/com/minimarket/UsuarioServiceImplTest.java")
    for item in [
        "Verifica que un usuario completo tenga nombre, apellido, email y direccion.",
        "Verifica que un usuario incompleto sea rechazado.",
        "Verifica que un usuario con rol permitido pueda ejecutar operaciones criticas.",
        "Verifica que un usuario con rol no autorizado no pueda registrar ventas.",
        "Verifica caminos de servicio como busqueda, guardado y eliminacion usando repositorio mockeado.",
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, "Dependencia simulada: UsuarioRepository, usando @Mock y @InjectMocks.")

    add_heading(doc, "Venta", 2)
    add_paragraph(doc, "Archivo de pruebas: src/test/java/com/minimarket/VentaServiceImplTest.java")
    for item in [
        "Registra una venta solo cuando el usuario existe, tiene datos completos, tiene rol autorizado y hay stock suficiente.",
        "Rechaza ventas sin stock suficiente.",
        "Rechaza ventas asociadas a usuarios con datos incompletos.",
        "Rechaza ventas de usuarios sin rol autorizado.",
        "Simula la consulta de stock por producto.",
        "Calcula correctamente el total sumando precio por cantidad en cada detalle.",
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, "Dependencias simuladas: VentaRepository, ProductoRepository y UsuarioRepository.")

    add_heading(doc, "Evidencia de ejecucion", 2)
    add_paragraph(doc, "Comando ejecutado:")
    add_code(doc, r".\mvnw.cmd verify")
    for item in [
        "Tests ejecutados: 17.",
        "Fallos: 0.",
        "Errores: 0.",
        "Omitidos: 0.",
        "JaCoCo: All coverage checks have been met.",
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, "Reporte generado: target/site/jacoco/index.html")
    add_coverage_table(doc)

    add_heading(doc, "4. Reflexion tecnica", 1)
    add_paragraph(
        doc,
        "Las pruebas unitarias mejoran la calidad del sistema porque permiten validar reglas criticas antes de llegar a produccion. "
        "En este caso, evitan errores operativos como ventas sin stock, ventas asociadas a usuarios inexistentes o incompletos, roles no autorizados y totales calculados de forma incorrecta.",
    )
    add_paragraph(
        doc,
        "Mockito permite simular dependencias entre ventas, usuarios y productos sin depender de una base de datos real. "
        "Esta estrategia hace que las pruebas sean rapidas, repetibles y enfocadas en la logica de negocio. "
        "Por ejemplo, para validar una venta sin stock, se configura ProductoRepository para devolver un producto con menor stock que la cantidad solicitada, "
        "y luego se verifica que el sistema lance una excepcion y no invoque ventaRepository.save.",
    )
    add_paragraph(
        doc,
        "Medir cobertura con JaCoCo aporta visibilidad sobre que partes del codigo estan siendo ejercitadas por las pruebas. "
        "En sistemas con multiples entidades relacionadas, esto ayuda a detectar reglas importantes sin prueba, reducir regresiones y establecer un criterio objetivo de calidad minima. "
        "La cobertura no reemplaza el buen diseno de casos, pero funciona como una barrera util para evitar que funcionalidades criticas queden sin verificacion automatizada.",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
